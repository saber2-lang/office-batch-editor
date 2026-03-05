import os
import re
import sys
import shutil
import datetime
from pathlib import Path

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    import openpyxl
    XLSX_SUPPORT = True
except ImportError:
    XLSX_SUPPORT = False

try:
    import xlrd
    import xlwt
    from xlutils.copy import copy as xl_copy
    XLS_SUPPORT = True
except ImportError:
    XLS_SUPPORT = False

try:
    import win32com.client
    import pywintypes
    DOC_SUPPORT = True
except ImportError:
    DOC_SUPPORT = False

SCRIPT_DIR = Path(__file__).parent.resolve()

# 跳过的临时文件前缀/后缀（Office 锁文件、Word 临时文件）
_SKIP_PREFIXES = ('~$', '~WRL')
_SKIP_SUFFIXES = ('.tmp',)

_scan_errors = []   # 每次扫描时重置，收集所有处理错误，最终汇总显示


# ══════════════════════════════════════════════
#  工具函数
# ══════════════════════════════════════════════

def print_banner():
    print("\n" + "=" * 50)
    print("       [批量内容替换工具]")
    print("=" * 50)
    print(f"  工作目录：{SCRIPT_DIR}")
    if not DOCX_SUPPORT:
        print("[警告] 未检测到 python-docx，Word(.docx)替换不可用")
        print("       安装命令：pip install python-docx")
    if not DOC_SUPPORT:
        print("[警告] 未检测到 pywin32，Word(.doc)替换不可用")
        print("       安装命令：pip install pywin32")
        print("       注意：需要本机已安装 Microsoft Word")
    if not XLSX_SUPPORT:
        print("[警告] 未检测到 openpyxl，Excel(.xlsx)替换不可用")
        print("       安装命令：pip install openpyxl")
    if not XLS_SUPPORT:
        print("[警告] 未检测到 xlrd/xlwt/xlutils，Excel(.xls)替换不可用")
        print("       安装命令：pip install xlrd==1.2.0 xlwt xlutils")
    print()


def get_input(prompt, allow_empty=False):
    while True:
        val = input(prompt).strip().strip('"').strip("'")
        if val or allow_empty:
            return val
        print("  [!] 输入不能为空，请重新输入\n")


def confirm(prompt):
    return input(prompt).strip().lower() == 'y'


def make_backup(target_dir):
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_dir = str(target_dir) + f"_backup_{timestamp}"
    shutil.copytree(target_dir, backup_dir)
    print(f"\n[OK] 备份已创建：{backup_dir}\n")
    return backup_dir


def _is_temp_file(path: Path) -> bool:
    """跳过 Office 产生的临时/锁文件"""
    name = path.name
    return (
        any(name.startswith(p) for p in _SKIP_PREFIXES)
        or any(name.endswith(s) for s in _SKIP_SUFFIXES)
    )


def _log_error(file_path, msg):
    """记录处理错误，最终在 print_stats 末尾汇总显示"""
    try:
        rel = Path(file_path).relative_to(SCRIPT_DIR)
    except ValueError:
        rel = Path(file_path)
    _scan_errors.append((str(rel), str(msg)))


def _get_skip_reason(path: Path) -> str:
    """
    检测文件是否为空文件或损坏文件，返回跳过原因字符串。
    空字符串表示文件正常，可以继续处理。
    """
    try:
        if path.stat().st_size == 0:
            return "空文件（0字节）"
        if path.suffix.lower() in ('.docx', '.xlsx'):
            import zipfile
            if not zipfile.is_zipfile(str(path)):
                return "非有效文档（文件损坏或格式异常）"
    except Exception:
        pass
    return ""


# ══════════════════════════════════════════════
#  Word 会话管理（单实例复用 + 断线重连）
# ══════════════════════════════════════════════

class WordSession:
    """
    在整个批处理期间维护一个 Word 进程实例。
    所有 .doc 文件共用同一个 Word Application，避免反复启动进程。
    若 Word 进程崩溃断连，自动重新启动并继续。
    """

    def __init__(self):
        self._app = None

    def _launch(self):
        # DispatchEx 始终创建独立新进程，不依附已有 Word 实例
        # 避免 "Property Visible can not be set" / 与用户打开的文档冲突
        app = win32com.client.DispatchEx("Word.Application")
        try:
            app.Visible = False
        except Exception:
            pass
        try:
            app.DisplayAlerts = False
        except Exception:
            pass
        self._app = app

    def _is_alive(self):
        try:
            _ = self._app.Version   # 轻量探针
            return True
        except Exception:
            return False

    def get(self):
        """获取可用的 Word Application，必要时自动重启"""
        if self._app is None:
            self._launch()
        elif not self._is_alive():
            self._quit_silent()
            self._launch()
        return self._app

    def _quit_silent(self):
        try:
            if self._app is not None:
                self._app.Quit()
        except Exception:
            pass
        self._app = None

    def quit(self):
        self._quit_silent()


# ══════════════════════════════════════════════
#  .doc 替换函数
# ══════════════════════════════════════════════

def replace_in_doc(file_path, rules, use_regex, preview_only,
                   word_session=None):
    """
    替换 .doc 文件（通过 win32com 调用 Word）
    rules: [(old_text, new_text), ...]，每个文件只打开一次，顺序应用所有规则。
    word_session: WordSession 实例（由 scan_and_replace 统一传入，复用 Word 进程）
    若未传入则本函数自己管理 Word 进程生命周期（兼容单文件调用）。
    """
    if not DOC_SUPPORT:
        return False

    own_session = word_session is None
    if own_session:
        word_session = WordSession()

    doc = None
    try:
        word = word_session.get()
        abs_path = str(file_path.resolve())
        doc = word.Documents.Open(abs_path, ReadOnly=False)

        # ── Undo 缓冲保护 ──────────────────────────────────────────────────────
        # COM Find/Replace 每次命中都写一条 Undo 记录，文件含大量匹配时 Undo 缓冲
        # 会急剧膨胀，触发"文件太大，无法保存"(-2146823160)，与文件实际大小无关。
        # UndoRecord.StartCustomRecord 将所有变更合并为单一 Undo 条目，大幅缩减占用。
        undo_started = False
        try:
            word.UndoRecord.StartCustomRecord("批量替换")
            undo_started = True
        except Exception:
            pass

        try:
            if use_regex:
                changed = _doc_replace_by_python_regex(doc, rules, preview_only)
            else:
                changed = _doc_replace_native(doc, rules, preview_only)
        finally:
            if undo_started:
                try:
                    word.UndoRecord.EndCustomRecord()
                except Exception:
                    pass

        if changed and not preview_only:
            # 保存前清空 Undo 栈，进一步释放内存
            try:
                doc.UndoClear()
            except Exception:
                pass
            # 禁用快速保存（增量格式）：强制完整写入，避免 .doc 文件体积虚胀
            try:
                word.Options.AllowFastSave = False
            except Exception:
                pass

            saved = False
            last_err = None

            # ── 保存方案1：直接 Save（最常见路径） ──────────────────────────
            try:
                doc.Save()
                saved = True
            except Exception as e1:
                last_err = e1

            # ── 保存方案2：SaveAs 同格式临时文件（绕过 Undo 缓冲残留） ──────
            if not saved:
                tmp_doc = abs_path + '.__tmp.doc'
                try:
                    doc.SaveAs(tmp_doc, FileFormat=0)  # wdFormatDocument=0
                    doc.Close(SaveChanges=False)
                    doc = None
                    Path(tmp_doc).replace(Path(abs_path))
                    saved = True
                except Exception as e2:
                    last_err = e2
                    try:
                        Path(tmp_doc).unlink(missing_ok=True)
                    except Exception:
                        pass

            # ── 保存方案3：SaveAs .docx → 重开 → SaveAs .doc ───────────────
            # .docx 使用完全不同的 XML 序列化路径，不受 Undo 缓冲限制，几乎不会失败。
            # 重新打开后另存为 .doc 即得到干净紧凑的文件。
            if not saved:
                tmp_docx = abs_path + '.__tmp.docx'
                doc2 = None
                try:
                    doc.SaveAs(tmp_docx, FileFormat=16)  # wdFormatXMLDocument=16
                    doc.Close(SaveChanges=False)
                    doc = None
                    doc2 = word_session.get().Documents.Open(tmp_docx, ReadOnly=False)
                    doc2.SaveAs(abs_path, FileFormat=0)
                    doc2.Close(SaveChanges=False)
                    doc2 = None
                    saved = True
                except Exception as e3:
                    last_err = e3
                finally:
                    if doc2 is not None:
                        try:
                            doc2.Close(SaveChanges=False)
                        except Exception:
                            pass
                    try:
                        Path(tmp_docx).unlink(missing_ok=True)
                    except Exception:
                        pass

            if not saved:
                raise RuntimeError(
                    f"三种保存方式均失败（UndoClear+Save / SaveAs.doc / docx转存）：{last_err}"
                )

        return bool(changed)

    except Exception as e:
        _log_error(file_path, e)
        return False

    finally:
        # 只关闭文档，不退出 Word（让 session 统一管理）
        try:
            if doc is not None:
                doc.Close(SaveChanges=False)
        except Exception:
            pass
        if own_session:
            word_session.quit()


def _word_find_replace(rng, old_text, new_text):
    """
    通用 Word Find/Replace，覆盖正文/脚注/文本框/页眉/页脚所有区域。

    关键点：
    ① 使用完整 11 个位置参数调用 Execute（关键字参数在某些文档 / 故事区域静默失败）
    ② Wrap=1（wdFindContinue）比 Wrap=0 更可靠，ReplaceAll 时不会造成无限循环
    参数顺序：
      FindText, MatchCase, MatchWholeWord, MatchWildcards,
      MatchSoundsLike, MatchAllWordForms, Forward,
      Wrap(1=wdFindContinue), Format, ReplaceWith, Replace(2=wdReplaceAll)
    """
    rng.Find.ClearFormatting()
    rng.Find.Replacement.ClearFormatting()
    rng.Find.Execute(old_text, True, False, False, False, False, True, 1, False, new_text, 2)


# 保持旧名兼容（两者行为完全相同）
_native_find_replace = _word_find_replace
_hf_find_replace     = _word_find_replace


def _fix_hf_tab_alignment(hf, section):
    """
    替换页眉/页脚文字后，若段落仍使用"大量空格"分隔左右两部分内容，
    自动将空格分隔符转换为右对齐制表位，防止替换后文字增多导致换行。

    检测模式：[左侧文字][3个以上空格][右侧文字]
    转为：    [左侧文字][\\t][右侧文字]  + 在右边距处设置右对齐制表位

    若段落已含制表符则跳过（说明已是制表位布局）。
    """
    try:
        ps = section.PageSetup
        avail_w = ps.PageWidth - ps.LeftMargin - ps.RightMargin
    except Exception:
        return

    for para in hf.Range.Paragraphs:
        try:
            txt = para.Range.Text or ''
            if '\t' in txt:
                continue
            m = re.search(r'^(.+?)(\s{3,})(.+\r?)$', txt, re.DOTALL)
            if not m:
                continue

            space_seq = m.group(2)
            para.Range.Find.ClearFormatting()
            para.Range.Find.Replacement.ClearFormatting()
            para.Range.Find.Execute(
                space_seq, True, False, False, False, False, True, 1, False, '\t', 2
            )
            try:
                fmt = para.Range.ParagraphFormat
                fmt.TabStops.ClearAll()
                fmt.TabStops.Add(Position=avail_w, Alignment=2, Leader=0)
            except Exception:
                pass
        except Exception:
            continue


def _doc_replace_native(doc, rules, preview_only):
    """
    非正则模式：分两路遍历文档所有区域，保留原始格式。
    rules: [(old_text, new_text), ...]，每个文档只打开一次，顺序应用所有规则。

    路线一（StoryRanges + 关键字参数 Execute）：
      正文(1) 脚注(2) 尾注(3) 批注(4) 文本框(5)

    路线二（Sections + 完整位置参数 Execute）：页眉 / 页脚
      页眉/页脚 Range 的 Find.Execute 需要完整位置参数 + Wrap=1，
      否则 Word COM 静默跳过，不报错也不替换。
      替换完成后自动修复"空格分隔"布局 → 右对齐制表位，防止文字增多换行。
    """
    changed = False

    # ── 路线一：正文 / 脚注 / 尾注 / 批注 / 文本框 ──
    for story_type in [1, 2, 3, 4, 5]:
        try:
            story = doc.StoryRanges(story_type)
        except Exception:
            continue

        while story is not None:
            try:
                next_story = story.NextStoryRange
            except Exception:
                next_story = None

            try:
                text = story.Text or ''
                for old, new in rules:
                    if old in text:
                        changed = True
                        if not preview_only:
                            _native_find_replace(story, old, new)
            except Exception:
                pass

            story = next_story

    # ── 路线二：页眉 / 页脚（完整位置参数 Execute + 制表位修复） ──
    for section in doc.Sections:
        for hf_collection in [section.Headers, section.Footers]:
            for hf_idx in [1, 2, 3]:
                try:
                    hf = hf_collection(hf_idx)
                    if not hf.Exists:
                        continue
                    rng = hf.Range
                    text = rng.Text or ''
                    hf_rule_changed = False
                    for old, new in rules:
                        if old in text:
                            changed = True
                            hf_rule_changed = True
                            if not preview_only:
                                _hf_find_replace(rng, old, new)
                    if hf_rule_changed and not preview_only:
                        _fix_hf_tab_alignment(hf, section)
                except Exception:
                    continue

    return changed


def _doc_replace_by_python_regex(doc, rules, preview_only):
    """
    正则模式：Python 逐段处理（正文 + 表格 + 页眉页脚）
    rules: [(old_pattern, new_text), ...]，顺序应用所有规则。
    """
    changed = False

    def process_range(rng):
        nonlocal changed
        try:
            text = rng.Text
            if not text:
                return
            new_val = text
            for old, new in rules:
                new_val = re.sub(old, new, new_val)
            if new_val != text:
                changed = True
                if not preview_only:
                    rng.Text = new_val
        except Exception:
            pass

    try:
        for para in doc.Paragraphs:
            process_range(para.Range)
    except Exception:
        pass

    try:
        for table in doc.Tables:
            for row in table.Rows:
                for cell in row.Cells:
                    for para in cell.Range.Paragraphs:
                        process_range(para.Range)
    except Exception:
        pass

    try:
        for section in doc.Sections:
            for hf_collection in [section.Headers, section.Footers]:
                for hf in hf_collection:
                    if hf.Exists:
                        process_range(hf.Range)
    except Exception:
        pass

    return changed


# ══════════════════════════════════════════════
#  .docx 替换函数
# ══════════════════════════════════════════════

def _iter_textbox_paragraphs(parent_elem):
    """
    迭代 XML 元素内所有文本框（w:txbxContent）中的段落。
    同时覆盖 DrawingML（wps:txbx）和 VML（v:textbox）两种文本框类型，
    二者均使用 w:txbxContent 存储段落内容。
    """
    try:
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph
        for txbx in parent_elem.iter(qn('w:txbxContent')):
            for p_elem in txbx.iter(qn('w:p')):
                yield Paragraph(p_elem, None)
    except Exception:
        return


_EMU_PER_TWIP = 635   # 914400 EMU/英寸 ÷ 1440 twips/英寸 = 635 EMU/twip


def _get_style_tab_stops(para):
    """
    遍历段落的样式继承链，收集所有继承的制表位。
    返回字典 {pos_twips: tab_type_str}，子样式优先（先遇到的不覆盖）。
    用于 _fix_hf_tab_alignment_docx 中清除居中等继承制表位，
    防止 \\t 被居中制表位拦截，无法到达右对齐制表位。
    """
    from docx.oxml.ns import qn
    tabs = {}
    style = para.style
    visited = set()
    while style is not None:
        sid = id(style)
        if sid in visited:
            break
        visited.add(sid)
        try:
            pPr = style.element.pPr
            if pPr is not None:
                tabs_elem = pPr.find(qn('w:tabs'))
                if tabs_elem is not None:
                    for tab in tabs_elem.findall(qn('w:tab')):
                        pos = int(tab.get(qn('w:pos'), 0))
                        val = tab.get(qn('w:val'), 'left')
                        if pos not in tabs:   # 子样式优先
                            tabs[pos] = val
        except Exception:
            pass
        try:
            style = style.base_style
        except Exception:
            break
    return tabs


def _fix_hf_tab_alignment_docx(section, hf):
    """
    .docx 页眉/页脚换行修复（等价于 .doc 的 _fix_hf_tab_alignment）。

    检测页眉/页脚段落中"大量空格分隔"布局：
        [左侧文字][≥3个空格][右侧文字]
    将空格替换为制表符，并设置右对齐制表位，
    使右侧内容始终贴靠页面右侧文字边界，防止替换后文字增多导致换行。

    关键修复：Word"Header"样式自带居中制表位（约 4513 twips），
    单个 \\t 会先被居中制表位拦截，导致右侧内容不贴右边距。
    解决方法：用 w:val="clear" 清除所有继承的非右对齐制表位，
    再显式设置右对齐制表位。
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    try:
        pw = int(section.page_width  or 0)
        lm = int(section.left_margin or 0)
        rm = int(section.right_margin or 0)
        avail_emu = pw - lm - rm
        if avail_emu <= 0:
            avail_emu = int(6 * 914400)   # 6 英寸 fallback
    except Exception:
        avail_emu = int(6 * 914400)

    avail_twips = max(1, avail_emu // _EMU_PER_TWIP)

    for para in hf.paragraphs:
        try:
            full_text = para.text          # python-docx 会合并所有 run 文本
            has_tab = '\t' in full_text

            if not has_tab:
                # 空格分隔模式：找到空格并替换为制表符
                m = re.search(r'^(.+?)(\s{3,})(.+)$', full_text)
                if not m:
                    continue

                space_seq = m.group(2)

                # ① 优先在单个 run 中找到空格序列并替换
                replaced = False
                for run in para.runs:
                    if space_seq in run.text:
                        run.text = run.text.replace(space_seq, '\t', 1)
                        replaced = True
                        break

                if not replaced:
                    # ② 空格跨多个 run：将所有 run 合并到第一个 run 后替换
                    if not para.runs:
                        continue
                    new_full = full_text.replace(space_seq, '\t', 1)
                    para.runs[0].text = new_full
                    for run in para.runs[1:]:
                        run.text = ''
            else:
                # 已含制表符：判断是否为"左\t右"分隔格式
                # 只处理恰好一个制表符且两侧均有文字的段落
                parts = full_text.split('\t', 1)
                if len(parts) != 2 or not parts[0].strip() or not parts[1].strip():
                    continue   # 格式复杂或不符合，跳过

            # ③ 读取样式继承链中的制表位，清除居中等非右对齐继承制表位
            #    防止 \t 被居中制表位（~4513 twips）拦截，无法到达右对齐制表位
            #    注意：即使段落已有 \t（上面 has_tab 分支），也需要更新制表位配置
            style_tabs = _get_style_tab_stops(para)

            pPr = para._p.get_or_add_pPr()
            for existing in pPr.findall(qn('w:tabs')):
                pPr.remove(existing)

            w_tabs = OxmlElement('w:tabs')

            # 为每个非右对齐、非 clear 的继承制表位添加 clear 标记，压制继承
            for pos, val in style_tabs.items():
                if val not in ('right', 'clear'):
                    clear_tab = OxmlElement('w:tab')
                    clear_tab.set(qn('w:val'), 'clear')
                    clear_tab.set(qn('w:pos'), str(pos))
                    w_tabs.append(clear_tab)

            # 设置右对齐制表位（始终使用页面计算值，确保贴近右边距）
            w_tab = OxmlElement('w:tab')
            w_tab.set(qn('w:val'), 'right')
            w_tab.set(qn('w:pos'), str(avail_twips))
            w_tabs.append(w_tab)

            pPr.append(w_tabs)

        except Exception:
            continue


def fix_docx_header_tabs(file_path):
    """
    对已修改但制表位未修复的 .docx 文件，单独执行页眉/页脚制表位修复。
    适用于早期替换工具运行后出现页眉换行的文档。
    """
    if not DOCX_SUPPORT:
        print("  [ERR] python-docx 未安装")
        return False
    try:
        doc = Document(file_path)
        for section in doc.sections:
            for hf in [
                section.header,            section.footer,
                section.first_page_header, section.first_page_footer,
                section.even_page_header,  section.even_page_footer,
            ]:
                if hf is not None:
                    _fix_hf_tab_alignment_docx(section, hf)
        doc.save(file_path)
        return True
    except Exception as e:
        print(f"  [ERR] 修复 {Path(file_path).name} 失败：{e}")
        return False


def replace_in_docx(file_path, rules, use_regex, preview_only):
    """
    替换 .docx 文件内容：正文、表格（含嵌套）、页眉页脚、文本框
    rules: [(old_text, new_text), ...]，每个文件只打开一次，顺序应用所有规则。
    格式保留策略：
      - 匹配在单个 run 内 → 100% 保留原格式
      - 匹配跨多个 run   → 保留首个 run 格式，并给出提示
    页眉/页脚若因替换导致文字增多，自动修复"空格分隔"布局为右对齐制表位。
    """
    if not DOCX_SUPPORT:
        return False
    try:
        doc = Document(file_path)
        changed = False

        def do_replace(text):
            for old, new in rules:
                text = re.sub(old, new, text) if use_regex else text.replace(old, new)
            return text

        def has_match(text):
            for old, _ in rules:
                if (bool(re.search(old, text)) if use_regex else (old in text)):
                    return True
            return False

        def replace_paragraph(para):
            nonlocal changed
            in_run_matched = False
            for run in para.runs:
                if not run.text:
                    continue
                replaced = do_replace(run.text)
                if replaced != run.text:
                    in_run_matched = True
                    changed = True
                    if not preview_only:
                        run.text = replaced
            if in_run_matched:
                return
            full_text = ''.join(run.text for run in para.runs)
            if not has_match(full_text):
                return
            changed = True
            if preview_only:
                return
            replaced_full = do_replace(full_text)
            print(f"    [跨Run] {file_path.name}：跨格式块匹配，替换区域继承首段格式")
            if para.runs:
                para.runs[0].text = replaced_full
                for run in para.runs[1:]:
                    run.text = ''

        def replace_paragraphs(paragraphs):
            for para in paragraphs:
                replace_paragraph(para)

        def replace_table(table):
            for row in table.rows:
                for cell in row.cells:
                    replace_paragraphs(cell.paragraphs)
                    for nested in cell.tables:
                        replace_table(nested)

        replace_paragraphs(doc.paragraphs)
        for table in doc.tables:
            replace_table(table)
        for section in doc.sections:
            for hf in [
                section.header,            section.footer,
                section.first_page_header, section.first_page_footer,
                section.even_page_header,  section.even_page_footer,
            ]:
                if hf is not None:
                    # 记录该 hf 在替换前是否含有搜索词（用于决定是否修复制表位）
                    hf_has_match = has_match(''.join(p.text for p in hf.paragraphs))
                    replace_paragraphs(hf.paragraphs)
                    for table in hf.tables:
                        replace_table(table)
                    # 若该页眉/页脚被改动，自动修复"空格分隔"布局防止换行
                    if hf_has_match and not preview_only:
                        _fix_hf_tab_alignment_docx(section, hf)
                    # 页眉/页脚内的文本框
                    try:
                        for para in _iter_textbox_paragraphs(hf._element):
                            replace_paragraph(para)
                    except Exception:
                        pass

        # 正文中的文本框（含表格单元格内的文本框）
        for para in _iter_textbox_paragraphs(doc.element.body):
            replace_paragraph(para)

        if changed and not preview_only:
            doc.save(file_path)
        return changed

    except Exception as e:
        _log_error(file_path, e)
        return False


# ══════════════════════════════════════════════
#  Excel 替换函数
# ══════════════════════════════════════════════

def replace_in_xlsx(file_path, rules, use_regex, preview_only):
    """替换 .xlsx 内容，单元格格式完全不受影响。rules: [(old, new), ...]"""
    if not XLSX_SUPPORT:
        return False
    try:
        wb = openpyxl.load_workbook(file_path)
        changed = False
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        new_val = cell.value
                        for old, new in rules:
                            new_val = re.sub(old, new, new_val) if use_regex \
                                      else new_val.replace(old, new)
                        if new_val != cell.value:
                            changed = True
                            if not preview_only:
                                cell.value = new_val
        if changed and not preview_only:
            wb.save(file_path)
        return changed
    except Exception as e:
        _log_error(file_path, e)
        return False


def replace_in_xls(file_path, rules, use_regex, preview_only):
    """替换 .xls 内容，通过 xlutils.copy 尽量保留格式。rules: [(old, new), ...]"""
    if not XLS_SUPPORT:
        return False
    try:
        rb = xlrd.open_workbook(str(file_path), formatting_info=True)
        changed = False

        def _apply_rules(val):
            for old, new in rules:
                val = re.sub(old, new, val) if use_regex else val.replace(old, new)
            return val

        # 第一阶段：快速检测是否有任何规则命中任何单元格
        outer_break = False
        for sheet in rb.sheets():
            if outer_break:
                break
            for row_idx in range(sheet.nrows):
                if outer_break:
                    break
                for col_idx in range(sheet.ncols):
                    cell = sheet.cell(row_idx, col_idx)
                    if cell.ctype == xlrd.XL_CELL_TEXT:
                        if _apply_rules(cell.value) != cell.value:
                            changed = True
                            outer_break = True
                            break

        if not changed or preview_only:
            return changed

        # 第二阶段：对所有单元格应用所有规则
        wb = xl_copy(rb)
        rb2 = xlrd.open_workbook(str(file_path), formatting_info=True)
        for sheet_idx, sheet in enumerate(rb2.sheets()):
            ws = wb.get_sheet(sheet_idx)
            for row_idx in range(sheet.nrows):
                for col_idx in range(sheet.ncols):
                    cell = sheet.cell(row_idx, col_idx)
                    if cell.ctype == xlrd.XL_CELL_TEXT:
                        new_val = _apply_rules(cell.value)
                        if new_val != cell.value:
                            ws.write(row_idx, col_idx, new_val)
        wb.save(str(file_path))
        return True
    except Exception as e:
        _log_error(file_path, e)
        return False


def replace_filename(path, rules, use_regex, preview_only):
    """
    对单个文件执行文件名替换，顺序应用所有规则。
    preview_only=True  → 只返回新名称，不实际重命名。
    preview_only=False → 执行重命名，失败时打印错误并返回 None。
    """
    name = path.name
    new_name = name
    for old, new in rules:
        new_name = re.sub(old, new, new_name) if use_regex else new_name.replace(old, new)
    if new_name == name:
        return None
    if not preview_only:
        try:
            path.rename(path.parent / new_name)
        except Exception as e:
            print(f"\n    [ERR] 文件重命名失败 {name} → {new_name}：{e}")
            return None
    return new_name


def replace_dirname(path, rules, use_regex, preview_only):
    """
    对单个目录执行重命名，顺序应用所有规则。
    preview_only=True  → 只返回新名称，不实际重命名。
    preview_only=False → 执行重命名，失败时打印错误并返回 None。
    """
    name = path.name
    new_name = name
    for old, new in rules:
        new_name = re.sub(old, new, new_name) if use_regex else new_name.replace(old, new)
    if new_name == name:
        return None
    if not preview_only:
        try:
            path.rename(path.parent / new_name)
        except Exception as e:
            print(f"\n    [ERR] 目录重命名失败 {name} → {new_name}：{e}")
            return None
    return new_name


# ══════════════════════════════════════════════
#  主扫描逻辑
# ══════════════════════════════════════════════

_SUPPORTED = {'.doc', '.docx', '.xlsx', '.xls'}


def _collect_supported_files():
    """收集工作目录下所有受支持类型的文件（跳过备份目录和临时文件）"""
    result = []
    for root, dirs, files in os.walk(SCRIPT_DIR, topdown=True):
        dirs[:] = [d for d in dirs if '_backup_' not in d]
        for f in files:
            fp = Path(root) / f
            if _is_temp_file(fp):
                continue
            if fp.suffix.lower() in _SUPPORTED:
                result.append(fp)
    return result


def _collect_directories():
    """
    收集工作目录下所有子目录（跳过备份目录），按路径深度降序排列。
    最深的目录排在最前，确保重命名时从内层向外层处理，
    防止父目录改名后子目录路径失效。
    """
    result = []
    for root, dirs, files in os.walk(SCRIPT_DIR, topdown=True):
        dirs[:] = [d for d in dirs if '_backup_' not in d]
        for d in dirs:
            result.append(Path(root) / d)
    result.sort(key=lambda p: len(p.parts), reverse=True)
    return result


def scan_and_replace(rules, options, preview_only=True, target_file=None):
    """
    rules    : [(old_text, new_text), ...] 替换规则列表，每个文件只打开一次应用所有规则
    options 字段：
      mode       : 'content'  → 仅替换文件内容（不改文件名/目录名）
                   'filename' → 仅替换文件名和目录名（不读内容，速度快）
                   'both'     → 同时替换文件内容、文件名和目录名
      use_regex  : bool
    target_file  : 若指定 Path 对象，则只处理该单个文件（忽略目录扫描）
    """
    global _scan_errors
    _scan_errors = []   # 每次扫描重置

    stats = {'content_changed': [], 'filename_changed': [], 'dirname_changed': []}
    mode = options.get('mode', 'content')
    do_content = mode in ('content', 'both')
    do_rename  = mode in ('filename', 'both')

    # 在做任何修改前统一收集路径（后续重命名不影响已收集的列表）
    if target_file is not None:
        # 单文件模式：只处理指定的一个文件，不扫描目录、不重命名目录
        all_supported = [target_file] if target_file.suffix.lower() in _SUPPORTED else []
        all_dirs = []
    else:
        all_supported = _collect_supported_files()
        all_dirs      = _collect_directories() if do_rename else []

    # ══ 文件内容替换 ══════════════════════════════
    if do_content:
        doc_files   = [f for f in all_supported if f.suffix.lower() == '.doc']
        other_files = [f for f in all_supported if f.suffix.lower() in _SUPPORTED - {'.doc'}]
        total = len(doc_files) + len(other_files)
        processed = 0

        ext_handler = {
            '.docx': replace_in_docx,
            '.xlsx': replace_in_xlsx,
            '.xls':  replace_in_xls,
        }

        word_session = WordSession() if (DOC_SUPPORT and doc_files) else None
        try:
            for file_path in doc_files:
                processed += 1
                # 预检：空文件或损坏文件直接跳过，不纳入报错
                skip_reason = _get_skip_reason(file_path)
                if skip_reason:
                    print(f"  [{processed}/{total}] {file_path.name} → 跳过（{skip_reason}）")
                    continue

                print(f"  [{processed}/{total}] {file_path.name}", end=' ', flush=True)
                err_before = len(_scan_errors)
                changed = replace_in_doc(
                    file_path, rules,
                    options['use_regex'], preview_only,
                    word_session=word_session
                )
                if len(_scan_errors) > err_before:
                    print("→ 出错（见末尾汇总）")
                elif changed:
                    print("→ 有变更")
                    stats['content_changed'].append((str(file_path), '.doc'))
                else:
                    print("→ 无匹配")
        finally:
            if word_session:
                word_session.quit()

        for file_path in other_files:
            ext = file_path.suffix.lower()
            processed += 1
            # 预检：空文件或损坏文件直接跳过，不纳入报错
            skip_reason = _get_skip_reason(file_path)
            if skip_reason:
                print(f"  [{processed}/{total}] {file_path.name} → 跳过（{skip_reason}）")
                continue

            print(f"  [{processed}/{total}] {file_path.name}", end=' ', flush=True)
            err_before = len(_scan_errors)
            changed = ext_handler[ext](
                file_path, rules, options['use_regex'], preview_only
            )
            if len(_scan_errors) > err_before:
                print("→ 出错（见末尾汇总）")
            elif changed:
                print("→ 有变更")
                stats['content_changed'].append((str(file_path), ext))
            else:
                print("→ 无匹配")

    # ══ 文件名 + 目录名替换 ═══════════════════════
    if do_rename:
        # filename-only 模式显示逐行进度；both 模式不输出（极快）
        show_progress = (mode == 'filename')

        # ── 文件名 ──
        n_files = len(all_supported)
        if show_progress and n_files:
            print(f"\n  --- 文件名 ({n_files} 个) ---")
        for i, file_path in enumerate(all_supported, 1):
            if show_progress:
                print(f"  [文件 {i}/{n_files}] {file_path.name}", end=' ', flush=True)
            new_name = replace_filename(
                file_path, rules, options['use_regex'], preview_only
            )
            if new_name:
                stats['filename_changed'].append((str(file_path), new_name))
                if show_progress:
                    print(f"→ {'命中' if preview_only else '已重命名'}: {new_name}")
            elif show_progress:
                print("→ 无匹配")

        # ── 目录名（从最深层开始，防止父目录改名导致路径失效）──
        n_dirs = len(all_dirs)
        if show_progress and n_dirs:
            print(f"\n  --- 目录名 ({n_dirs} 个) ---")
        for i, dir_path in enumerate(all_dirs, 1):
            if show_progress:
                print(f"  [目录 {i}/{n_dirs}] {dir_path.name}", end=' ', flush=True)
            new_name = replace_dirname(
                dir_path, rules, options['use_regex'], preview_only
            )
            if new_name:
                stats['dirname_changed'].append((str(dir_path), new_name))
                if show_progress:
                    print(f"→ {'命中' if preview_only else '已重命名'}: {new_name}")
            elif show_progress:
                print("→ 无匹配")

    stats['errors'] = list(_scan_errors)
    return stats


def print_stats(stats, preview_only, mode):
    label = "[预览结果]" if preview_only else "[执行结果]"
    print(f"\n{'-'*50}")
    print(f"  {label}")
    print(f"{'-'*50}")

    if mode in ('content', 'both'):
        if stats['content_changed']:
            print(f"\n[文件内容] 将被修改 {len(stats['content_changed'])} 个：")
            for path, ext in stats['content_changed']:
                print(f"   - {Path(path).name}  ({ext})")
        else:
            print("\n[文件内容] 无匹配项，无需修改")

    if mode in ('filename', 'both'):
        verb = "命中" if preview_only else "已重命名"
        if stats['filename_changed']:
            print(f"\n[文件名]  {verb} {len(stats['filename_changed'])} 个：")
            for old, new in stats['filename_changed']:
                print(f"   - {Path(old).name}  →  {new}")
        else:
            print("\n[文件名]  无匹配项，无需修改")

        if stats['dirname_changed']:
            print(f"\n[目录名]  {verb} {len(stats['dirname_changed'])} 个：")
            for old, new in stats['dirname_changed']:
                print(f"   - {Path(old).name}  →  {new}")
        else:
            print("\n[目录名]  无匹配项，无需修改")

    total = (len(stats['content_changed'])
             + len(stats['filename_changed'])
             + len(stats['dirname_changed']))
    print(f"\n{'-'*50}")
    print(f"  共计影响项目：{total} 个")
    print(f"{'-'*50}\n")

    # 末尾汇总显示错误（不包含空文件跳过，仅处理失败的文件）
    errors = stats.get('errors', [])
    if errors:
        print(f"{'='*50}")
        print(f"  [错误汇总] 处理失败 {len(errors)} 个文件：")
        print(f"{'='*50}")
        for rel_path, msg in errors:
            print(f"   - {rel_path}")
            print(f"     原因：{msg}")
        print(f"{'='*50}\n")


# ══════════════════════════════════════════════
#  主入口
# ══════════════════════════════════════════════

def _build_arg_parser():
    """构建命令行参数解析器（独立函数方便单元测试）"""
    import argparse
    parser = argparse.ArgumentParser(
        prog='rename_tool',
        description='批量文件内容 / 文件名替换工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "使用示例：\n"
            "  python rename_tool.py\n"
            "      交互模式，扫描脚本所在目录\n\n"
            "  python rename_tool.py --dir \"C:\\\\体系文件\"\n"
            "      指定目标目录（交互模式输入其余参数）\n\n"
            "  python rename_tool.py --file \"doc.docx\"\n"
            "      只处理指定的单个文件\n\n"
            "  python rename_tool.py --old \"旧公司名\" --new \"新公司名\" --mode 1 --yes --no-backup\n"
            "      全非交互模式，替换脚本目录下所有文件内容\n\n"
            "  python rename_tool.py --dir \"C:\\\\体系文件\" --old \"NW-\" --new \"KSMD-\" --mode 3 --yes\n"
            "      指定目录，同时替换内容和文件名，不询问确认"
        ),
    )
    parser.add_argument('--dir',  '-d', metavar='PATH',
                        help='目标目录（默认：脚本所在目录）')
    parser.add_argument('--file', '-f', metavar='PATH',
                        help='仅处理指定的单个文件')
    parser.add_argument('--old',  '-o', metavar='TEXT', action='append',
                        help='要查找的旧内容（可多次指定，与 --new 一一对应）')
    parser.add_argument('--new',  '-n', metavar='TEXT', action='append',
                        help='替换后的新内容（可多次指定，与 --old 一一对应，默认空字符串）')
    parser.add_argument('--mode', '-m', metavar='{1,2,3}',
                        choices=['1', '2', '3'],
                        help='1=仅替换文件内容  2=仅替换文件名/目录名  3=两者均替换')
    parser.add_argument('--regex', '-r', action='store_true',
                        help='使用正则表达式模式')
    parser.add_argument('--no-backup', action='store_true',
                        help='不创建备份（谨慎使用）')
    parser.add_argument('--yes', '-y', action='store_true',
                        help='跳过所有确认直接执行（需配合 --old 使用）')
    return parser


def main(argv=None):
    """
    argv=None  → 从 sys.argv 读取（正常命令行调用）
    argv=[...] → 使用指定列表（方便测试）
    """
    args = _build_arg_parser().parse_args(argv)

    print_banner()

    # ── 确定工作目录 / 目标文件 ──────────────────────────────────────
    global SCRIPT_DIR
    target_file = None

    if args.file and args.dir:
        print("[!] --file 和 --dir 不能同时使用，请只指定其中一个")
        sys.exit(1)

    if args.file:
        target_file = Path(args.file).resolve()
        if not target_file.exists():
            print(f"[!] 文件不存在：{target_file}")
            sys.exit(1)
        if target_file.suffix.lower() not in _SUPPORTED:
            print(f"[!] 不支持的文件类型：{target_file.suffix}"
                  f"（支持：{', '.join(sorted(_SUPPORTED))}）")
            sys.exit(1)
        SCRIPT_DIR = target_file.parent
        print(f"  目标文件：{target_file.name}")
    elif args.dir:
        SCRIPT_DIR = Path(args.dir).resolve()
        if not SCRIPT_DIR.is_dir():
            print(f"[!] 目录不存在：{SCRIPT_DIR}")
            sys.exit(1)
        print(f"  目标目录：{SCRIPT_DIR}")

    # ── 获取查找 / 替换内容（支持多条规则）────────────────────────────
    if args.old is not None:
        # CLI 模式：--old/--new 可多次指定，一一对应
        olds = args.old
        news = args.new if args.new else [''] * len(olds)
        if len(olds) != len(news):
            print(f"[!] --old 和 --new 数量不一致（{len(olds)} 对 {len(news)}），请一一对应")
            sys.exit(1)
        rules = list(zip(olds, news))
        if len(rules) == 1:
            print(f"  查找内容：{rules[0][0]!r}")
            print(f"  替换内容：{rules[0][1]!r}")
        else:
            print(f"  替换规则（{len(rules)} 条）：")
            for i, (old, new) in enumerate(rules, 1):
                print(f"    {i}. {old!r}  →  {new!r}")
    else:
        rules = []
        rule_idx = 1
        while True:
            if rule_idx > 1:
                print()
            prefix = f"[规则 {rule_idx}] " if rule_idx > 1 else ""
            old = get_input(f"{prefix}请输入要查找的内容（旧内容）：\n> ")
            new = get_input(f"{prefix}请输入替换后的内容（新内容，可为空）：\n> ", allow_empty=True)
            rules.append((old, new))
            if not confirm("是否继续添加替换规则？(y/n): "):
                break
            rule_idx += 1

    # ── 正则模式 ─────────────────────────────────────────────────────
    use_regex = args.regex
    if not use_regex and args.old is None:
        print()
        use_regex = confirm("是否使用正则表达式？(y/n，普通用户选 n): ")
    if use_regex:
        all_valid = True
        for old, _ in rules:
            try:
                re.compile(old)
            except re.error as e:
                print(f"  [!] 规则 {old!r} 正则有误：{e}，已自动切换为普通模式")
                use_regex = False
                all_valid = False
                break
        if all_valid:
            print("  [OK] 正则表达式验证通过")

    # ── 操作模式 ─────────────────────────────────────────────────────
    if args.mode:
        mode_choice = args.mode
        _mode_label = {'1': '仅内容', '2': '仅文件名/目录名', '3': '内容+文件名'}
        print(f"  操作模式：{_mode_label[mode_choice]}")
    elif target_file:
        mode_choice = '1'   # 单文件默认只替换内容
    else:
        print()
        print("请选择操作模式：")
        print("  1. 仅替换文件内容（不修改文件名）")
        print("  2. 仅替换文件名  （不读取文件内容，速度快）")
        print("  3. 同时替换文件内容和文件名")
        while True:
            choice = input("> ").strip()
            if choice in ('1', '2', '3'):
                break
            print("  [!] 请输入 1、2 或 3")
        mode_choice = choice

    mode = {'1': 'content', '2': 'filename', '3': 'both'}[mode_choice]
    options = {'mode': mode, 'use_regex': use_regex}

    # ── 预览 ─────────────────────────────────────────────────────────
    print("\n正在扫描，生成预览...")
    preview_stats = scan_and_replace(
        rules, options, preview_only=True, target_file=target_file
    )
    print_stats(preview_stats, preview_only=True, mode=mode)

    total = (len(preview_stats['content_changed'])
             + len(preview_stats['filename_changed'])
             + len(preview_stats['dirname_changed']))
    if total == 0:
        print("未找到任何匹配内容，程序退出。")
        if not args.yes:
            input("按回车键退出...")
        return

    # ── 备份 ─────────────────────────────────────────────────────────
    if not args.no_backup:
        # --yes 时自动备份；否则询问
        do_backup = args.yes or confirm("执行前是否自动备份整个文件夹？（强烈建议 y）: ")
        if do_backup:
            make_backup(SCRIPT_DIR)

    # ── 确认执行 ─────────────────────────────────────────────────────
    if not args.yes:
        if not confirm("确认执行以上替换操作？(y/n): "):
            print("\n已取消，未做任何修改。")
            input("按回车键退出...")
            return

    # ── 执行 ─────────────────────────────────────────────────────────
    print("\n正在执行替换...")
    real_stats = scan_and_replace(
        rules, options, preview_only=False, target_file=target_file
    )
    print_stats(real_stats, preview_only=False, mode=mode)
    print("全部替换完成！")
    if not args.yes:
        input("按回车键退出...")


if __name__ == "__main__":
    try:
        main()
    except SystemExit:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        input("发生异常，按回车键退出...")

